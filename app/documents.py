"""Best-effort plain-text extraction from common document formats, so attached
PDF / Word / Excel / CSV files become readable context for the model instead of
opaque "binary" blobs.

Each format's parser is imported lazily: if its optional dependency
(``pypdf`` / ``python-docx`` / ``openpyxl``) is not installed, extraction simply
returns ``None`` and the caller falls back to the previous behaviour rather than
crashing. CSV/TSV use the standard library and always work.
"""

from __future__ import annotations

import csv as _csv
from pathlib import Path

# Max characters of extracted text to keep — mirrors MAX_TEXT_FILE_BYTES so a huge
# document can't blow up the prompt / context window.
MAX_CHARS = 80_000

DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xlsm", ".csv", ".tsv"}


def extract_text(path: Path) -> str | None:
    """Return extracted plain text for a supported document, or None when the
    format is unsupported, the optional parser is missing, or parsing fails."""
    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            return _truncate(_extract_pdf(path))
        if suffix == ".docx":
            return _truncate(_extract_docx(path))
        if suffix in (".xlsx", ".xlsm"):
            return _truncate(_extract_xlsx(path))
        if suffix in (".csv", ".tsv"):
            return _truncate(_extract_csv(path, "\t" if suffix == ".tsv" else ","))
    except Exception:
        return None
    return None


def extractor_available(suffix: str) -> bool:
    """Whether the optional dependency needed for ``suffix`` is importable."""
    suffix = suffix.lower()
    try:
        if suffix == ".pdf":
            import pypdf  # noqa: F401
        elif suffix == ".docx":
            import docx  # noqa: F401
        elif suffix in (".xlsx", ".xlsm"):
            import openpyxl  # noqa: F401
        elif suffix in (".csv", ".tsv"):
            return True
        else:
            return False
        return True
    except Exception:
        return False


def _truncate(text: str | None) -> str | None:
    if not text or not text.strip():
        return None
    text = text.strip()
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n... (document truncated)"
    return text


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join(p for p in pages if p)


def _extract_docx(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(path: Path) -> str:
    import openpyxl

    workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    out: list[str] = []
    for sheet in workbook.worksheets:
        out.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if value is None else str(value) for value in row]
            if any(cells):
                out.append(", ".join(cells))
    return "\n".join(out)


def _extract_csv(path: Path, delimiter: str) -> str:
    with open(path, newline="", encoding="utf-8", errors="replace") as handle:
        rows = list(_csv.reader(handle, delimiter=delimiter))
    return "\n".join(", ".join(row) for row in rows)
